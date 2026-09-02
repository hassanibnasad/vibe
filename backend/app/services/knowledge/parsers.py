"""
Document parsers for the knowledge ingestion pipeline.

Supports: Markdown (.md), plain text (.txt), PDF (.pdf), DOCX (.docx),
and raw bytes with MIME sniffing.

Each parser returns a ``ParsedDocument`` with ``title``, ``raw_text``,
and ``metadata``.  No LLM calls happen here — parsing is pure CPU / IO.
"""

from __future__ import annotations

import io
import mimetypes
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import structlog

logger = structlog.get_logger()


@dataclass
class ParsedDocument:
    """Output of any document parser."""

    title: str
    raw_text: str
    metadata: dict[str, Any] = field(default_factory=dict)


class MarkdownParser:
    """Parse Markdown / plain-text files."""

    def parse_path(self, path: Path) -> ParsedDocument:
        raw = path.read_text(encoding="utf-8", errors="replace")
        title = self._extract_title(raw) or path.stem.replace("-", " ").replace("_", " ").title()
        return ParsedDocument(
            title=title,
            raw_text=raw,
            metadata={"source_file": str(path), "extension": path.suffix},
        )

    def parse_bytes(self, content: bytes, filename: str = "") -> ParsedDocument:
        raw = content.decode("utf-8", errors="replace")
        title = self._extract_title(raw) or Path(filename).stem.replace("-", " ").title()
        return ParsedDocument(
            title=title,
            raw_text=raw,
            metadata={"filename": filename, "extension": Path(filename).suffix},
        )

    @staticmethod
    def _extract_title(text: str) -> str:
        """Return the first H1 heading if present, else empty string."""
        for line in text.splitlines():
            stripped = line.strip()
            if stripped.startswith("# "):
                return stripped[2:].strip()
        return ""


class PDFParser:
    """Parse PDF files using pypdf (pure-Python, no native deps)."""

    def parse_path(self, path: Path) -> ParsedDocument:
        return self._parse(path.read_bytes(), filename=path.name)

    def parse_bytes(self, content: bytes, filename: str = "") -> ParsedDocument:
        return self._parse(content, filename=filename)

    @staticmethod
    def _parse(content: bytes, filename: str) -> ParsedDocument:
        try:
            import pypdf  # lazy import — optional dep

            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            raw_text = "\n\n".join(pages)

            info = reader.metadata or {}
            title = str(info.get("/Title", "")).strip() or Path(filename).stem.replace("-", " ").title()

            return ParsedDocument(
                title=title,
                raw_text=raw_text,
                metadata={
                    "filename": filename,
                    "page_count": len(reader.pages),
                    "pdf_author": str(info.get("/Author", "")),
                },
            )
        except ImportError:
            logger.warning("pypdf_not_installed", hint="pip install pypdf to enable PDF ingestion")
            raise RuntimeError(
                "pypdf is required for PDF ingestion. Install it with: pip install pypdf"
            )


class DocxParser:
    """Parse DOCX files using python-docx."""

    def parse_path(self, path: Path) -> ParsedDocument:
        return self._parse(path.read_bytes(), filename=path.name)

    def parse_bytes(self, content: bytes, filename: str = "") -> ParsedDocument:
        return self._parse(content, filename=filename)

    @staticmethod
    def _parse(content: bytes, filename: str) -> ParsedDocument:
        try:
            import docx  # lazy import — optional dep

            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            raw_text = "\n\n".join(paragraphs)
            title = paragraphs[0].strip() if paragraphs else Path(filename).stem.replace("-", " ").title()

            return ParsedDocument(
                title=title,
                raw_text=raw_text,
                metadata={"filename": filename},
            )
        except ImportError:
            logger.warning("python-docx_not_installed", hint="pip install python-docx to enable DOCX ingestion")
            raise RuntimeError(
                "python-docx is required for DOCX ingestion. Install it with: pip install python-docx"
            )


# ── MIME → parser routing ────────────────────────────────────────────────────

_MIME_PARSERS: dict[str, type] = {
    "text/markdown": MarkdownParser,
    "text/x-markdown": MarkdownParser,
    "text/plain": MarkdownParser,
    "application/pdf": PDFParser,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocxParser,
}

_EXT_PARSERS: dict[str, type] = {
    ".md": MarkdownParser,
    ".txt": MarkdownParser,
    ".pdf": PDFParser,
    ".docx": DocxParser,
}

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset(_EXT_PARSERS.keys())
SUPPORTED_MIME_TYPES: frozenset[str] = frozenset(_MIME_PARSERS.keys())


class DocumentParser:
    """
    Unified entry-point.  Routes to the correct sub-parser based on file
    extension (``parse_path``) or MIME type (``parse_bytes``).
    """

    def parse_path(self, path: Path) -> ParsedDocument:
        """Parse a file from disk.  Raises ``ValueError`` for unsupported extensions."""
        ext = path.suffix.lower()
        parser_cls = _EXT_PARSERS.get(ext)
        if parser_cls is None:
            raise ValueError(
                f"Unsupported file extension '{ext}'. "
                f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
            )
        logger.debug("document_parser.parse_path", path=str(path), parser=parser_cls.__name__)
        return parser_cls().parse_path(path)

    def parse_bytes(
        self,
        content: bytes,
        filename: str,
        mime_type: str | None = None,
    ) -> ParsedDocument:
        """
        Parse raw bytes.  Resolution order:
        1. Explicit ``mime_type`` argument
        2. Guess from ``filename`` extension
        3. Fallback to MarkdownParser for plain text
        """
        resolved_mime = mime_type
        if not resolved_mime and filename:
            resolved_mime, _ = mimetypes.guess_type(filename)

        parser_cls = _MIME_PARSERS.get(resolved_mime or "") or _EXT_PARSERS.get(
            Path(filename).suffix.lower(), MarkdownParser
        )
        logger.debug(
            "document_parser.parse_bytes",
            filename=filename,
            mime=resolved_mime,
            parser=parser_cls.__name__,
        )
        return parser_cls().parse_bytes(content, filename=filename)
